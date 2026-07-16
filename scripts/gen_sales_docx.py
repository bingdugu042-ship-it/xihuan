# -*- coding: utf-8 -*-
"""生成《星屿市定价与销售策略》Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_cn(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True)
    return p


def add_para(text, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(title.add_run("星屿市 / 星屹游乐园"), size=22, bold=True, color=RGBColor(0x1F, 0x7A, 0x94))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(sub.add_run("定价方案 · 成团销售 · NFC 钥匙包双重协同策略"), size=14, bold=True)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(meta.add_run("内部销售策略草案 · 2026年7月"), size=10, color=RGBColor(0x5A, 0x6F, 0x7A))

add_para(
    "一句话结论：当前定制价整体偏低、适合开荒吸粉；NFC+谷子应以「入园钥匙包」绑定数字完整解锁销售，"
    "消化库存的同时抬高客单价。散谷单独卖难动，绑钥匙才有故事。"
)

# 一
add_heading_cn("一、定制定价判断与建议", 1)
add_heading_cn("1.1 原方案高低判断", 2)
add_table(
    ["项目", "原报价", "评价"],
    [
        ["单角色", "30 元", "偏低～刚好，引流优秀"],
        ["多角色", "50 元", "偏低"],
        ["每加一个功能", "+10 元", "偏低，建议 +15"],
        ["超大型 22 玩法", "100 元", "明显过低（工程量不对等）"],
    ],
)
add_para("总评：不是定高了，是定低了。低价适合开荒测盘（建议限时+限量），不适合当永久标准价。")

add_heading_cn("1.2 推荐落地版（微涨仍好成交）", 2)
add_table(
    ["档位", "建议价", "说明"],
    [
        ["单角色定制", "39 元", "含人设+开场+基础指令，含一轮修改"],
        ["多角色（2–3人）", "69 元", "群戏/多接星者包"],
        ["每加 1 个系统功能", "+15 元", "身体面板联动、盖章轨、硬指引阶段、生图管线等"],
        ["22 设施玩法包级", "299–399 元起", "不是 100；100 只适合「现成品小改」"],
    ],
)
add_para(
    "若坚持活动价 30/50/+10：写明「活动至×月」，并限制周接单量（如每周≤3单）。"
    "100 元只能覆盖「现有引擎换皮/少量子区域」，不能从零定做 22 玩法。"
)

add_heading_cn("1.3 定制交付边界（防亏本）", 2)
add_para("含：人设 JSON/卡面、开场白、设施对接说明、一轮修改。")
add_para("不含：自备 API Key、长期陪聊调 prompt、商用授权、无限修改。")
add_para("加价项：立绘管线、语音、多设施关卡文案、按月维护。")

# 二
add_heading_cn("二、开团（团购）三套方案", 1)
add_para("原则：成团解决「半成品→完整可玩」；成团后另设常驻价；不要二开更低价撬一团。")

add_heading_cn("2.1 方案 A · 低价破冰（推荐首发）", 2)
add_table(
    ["项", "内容"],
    [
        ["成团门槛", "30 人"],
        ["成团价", "19.9 元/人"],
        ["成团后常驻价", "29.9 元"],
        ["交付", "半成品 → 完整可玩版（全设施入园、盖章/档案/硬指引等补全）"],
        ["约入账", "30×19.9 ≈ 597 元 + 团后零售"],
    ],
)

add_heading_cn("2.2 方案 B · 平衡档（长期主推）", 2)
add_table(
    ["项", "内容"],
    [
        ["成团门槛", "50 人"],
        ["成团价", "29.9 元/人"],
        ["早鸟（前20人）", "24.9 元"],
        ["成团后常驻价", "45–49 元"],
        ["约入账", "50×29.9 ≈ 1495 元"],
    ],
)

add_heading_cn("2.3 方案 C · 精品慢卖", 2)
add_table(
    ["项", "内容"],
    [
        ["成团门槛", "20 人"],
        ["成团价", "49 元/人"],
        ["成团后常驻价", "68–79 元"],
        ["附加", "一年内小更新 / 优先反馈通道"],
        ["约入账", "20×49 = 980 元"],
    ],
)

# 三
add_heading_cn("三、NFC 启动卡 + 谷子：杀器定位", 1)
add_para("核心打法：NFC 卡不是「周边主商品」，而是「入园钥匙」。")
add_para("实体卡 = 身份 / 入场凭证 / 成团纪念")
add_para("数字半成品免费 = 试用")
add_para("成团解锁完整包 = 内容付费")
add_para("谷子/打印小卡 = 钥匙配菜（消化库存）")
add_para("统一话术：免费先玩；成团解锁完整星屿；NFC 启动卡是实体入园证，开盒送配套小谷。")

add_heading_cn("3.1 为什么谷子单独不好卖", 2)
add_para("1）没有「必须用卡/必须进园」的故事，买家当普通谷；")
add_para("2）缺 SKU 打包与稀缺，堆货≠货盘；")
add_para(
    "3）改造：NFC 绑定护照编号或设施章皮肤；小卡印 22 章/接星者，对接游戏集邮；"
    "包装做成「扫卡进园」仪式（技术可为 URL+邀请码/解锁码）。"
)

# 四
add_heading_cn("四、双重协同定价表", 1)

add_heading_cn("4.1 数字 alone（无实体）", 2)
add_table(
    ["档", "价格", "给什么"],
    [
        ["成团", "19.9（满30人）", "完整可玩包"],
        ["团后常驻", "29.9", "同上"],
        ["定制", "活动 30/50/+10；22设施级 299 起", "定制内容"],
    ],
)

add_heading_cn("4.2 实体钥匙包（主推杀器）", 2)
add_table(
    ["档", "价格", "给什么"],
    [
        ["入园钥匙包", "49–59 元（建议定 59）", "NFC×1 + 小卡/谷配套 + 赠数字完整解锁 1 份"],
        ["典藏钥匙包", "89–99 元", "卡 + 大件谷 + 数字完整包 + 限定身份码/专属开场"],
        ["钥匙+定制", "钥匙价+定制×0.9", "例：49+30→套装 69"],
    ],
)
add_para("要点：NFC 必须绑定「数字完整解锁」，否则卡退回普通谷子。谷子不要标高价单卖；标成「开盒附赠价值 xx，现随卡打包」。")

add_heading_cn("4.3 库存消化", 2)
add_table(
    ["卖法", "价格", "目的"],
    [
        ["散谷单独清仓", "9.9–19.9", "甩货，不冲击主价"],
        ["绑进钥匙包", "心里摊 15–25 成本", "出库存+抬客单"],
        ["满 3 件谷", "29.9", "仅清仓场/直播"],
    ],
)

add_heading_cn("4.4 对外统一价目表", 2)
add_table(
    ["买什么", "钱", "数字", "实体"],
    [
        ["白嫖体验", "0", "半成品", "无"],
        ["成团完整版", "19.9", "完整包", "无"],
        ["入园钥匙包", "59", "完整包", "NFC+谷"],
        ["典藏", "99", "完整+小特权", "大谷"],
        ["定制", "另计", "人设/功能", "可加 9.9 印小卡"],
    ],
)

add_heading_cn("4.5 交叉优惠（只给付过钱的人）", 2)
add_para("已成团再补钥匙：补差价约 30 元。")
add_para("已买钥匙再定制：定制 9 折。")
add_para("已定制再补卡：卡包减 10 元。")
add_para("形成数字与实体互相追着买，而不是两条互拆价。")

# 五
add_heading_cn("五、销售漏斗", 1)
add_para("半成品免费玩 → 想要完整版/仪式感 → 分流：")
add_para("　A. 成团 19.9　→　可加购钥匙包补差价")
add_para("　B. NFC 钥匙包 59　→　自动含完整数字　→　加价定制　→　复购第二张身份卡/典藏")
add_para("执行原则：")
add_para("1）免费层玩到「差点意思」，文案写清完整版成团解锁；")
add_para("2）成团层便宜冲人数，成功后发码；")
add_para("3）钥匙层更贵，货值靠实体+附赠解锁；")
add_para("4）定制层给已付费用户折扣，培养回头。")

# 六
add_heading_cn("六、两套开团脚本", 1)
add_heading_cn("6.1 开团一：纯数字破冰", 2)
add_para("满 30 人 × 19.9。赠：电子护照壁纸 / 数字小卡 PDF（不动实体库存）。")
add_heading_cn("6.2 开团二：钥匙成团（消化库存）", 2)
add_para("满 20 人 × 59。必含：NFC + 配套小卡/谷 + 数字完整解锁。限量：仅限现货库存 N 份。")
add_para("成团后：数字 alone 29.9；钥匙包常驻 69（略高于成团）。写死：二开不低于一团价。")

# 七
add_heading_cn("七、默认推荐组合（直接定稿）", 1)
add_para("成团破冰：30 人 × 19.9，团后 29.9。", bold=True)
add_para("杀器主售：入园钥匙包 59（NFC+赠谷+含完整数字）。", bold=True)
add_para("典藏：99（限量，专治最好看的一批囤货）。", bold=True)
add_para("定制：活动期可用 30/50/+10，但「22 玩法」改为 299 起，标注活动截止日。", bold=True)
add_para("2 个月后：定制涨至 39/69/+15，成团切方案 B。")
add_para("散谷：9.9–19.9 清仓，不占主 C 位。")

# 八
add_heading_cn("八、七天执行清单", 1)
add_para("1. 盘库存：NFC 卡数、可打包成「钥匙包」的份数。")
add_para("2. 做入园钥匙包开箱图/短视频：扫卡 → 出邀请码 → 进完整版。")
add_para("3. 并行开两个团：数字 19.9 / 钥匙 59。")
add_para("4. 免费半成品页顶栏：成团进度 + 钥匙购买入口。")
add_para("5. 每周一次清仓场：只卖散谷 9.9，不冲钥匙主线。")

# 九
add_heading_cn("九、对外一句总话术", 1)
add_para(
    "免费体验半成品；成团解锁完整包；NFC 启动卡是实体入园证（开盒送配套谷子）；角色/功能定制另计。"
)
add_para("")
add_para("—— 文档结束。可按实际库存把「钥匙包份数 N、典藏清单」填进附表后对外使用。", size=10)

out = r"C:\Users\wenwangcao\Desktop\星屿市_定价与销售策略.docx"
doc.save(out)
print("saved:", out)
